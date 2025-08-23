"""
HR Automation - Word Document Generator
Creates formatted Word documents with shortlisted candidate information
"""

import json
from typing import List, Dict, Any, Union
from datetime import datetime
import logging
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    Document = None  # Define as None when not available
    logging.warning("python-docx not installed. Word document generation will be limited.")

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

class CandidateDocumentGenerator:
    """Generates Word documents for shortlisted candidates"""
    
    def __init__(self):
        if not DOCX_AVAILABLE:
            logger.warning("Word document functionality requires python-docx. Install with: pip install python-docx")
    
    def create_shortlist_document(self, shortlists: Dict[str, List[Dict[str, Any]]], 
                                output_filename: str = None) -> str:
        """
        Create a comprehensive Word document with all shortlists
        
        Args:
            shortlists: Dictionary of job titles and their candidate matches
            output_filename: Optional custom filename
            
        Returns:
            Path to the generated document
        """
        if not DOCX_AVAILABLE:
            # Fallback to text-based document
            return self._create_text_document(shortlists, output_filename)
        
        # Generate filename if not provided
        if not output_filename:
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            output_filename = f"candidate_shortlists_{timestamp}.docx"
        
        # Create Word document
        doc = Document()
        
        # Add title and header
        self._add_document_header(doc)
        
        # Add shortlists for each job
        for job_title, candidates in shortlists.items():
            self._add_job_section(doc, job_title, candidates)
        
        # Add footer
        self._add_document_footer(doc)
        
        # Save document
        doc.save(output_filename)
        logger.info(f"Shortlist document saved as: {output_filename}")
        
        return output_filename
    
    def create_selected_candidates_document(self, selected_candidates: List[Dict[str, Any]], 
                                          job_title: str = "Selected Candidates",
                                          output_filename: str = None) -> str:
        """
        Create a Word document for specifically selected candidates
        
        Args:
            selected_candidates: List of selected candidate dictionaries
            job_title: Job title for the document
            output_filename: Optional custom filename
            
        Returns:
            Path to the generated document
        """
        if not DOCX_AVAILABLE:
            return self._create_text_document_selected(selected_candidates, job_title, output_filename)
        
        if not output_filename:
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            safe_job_title = "".join(c for c in job_title if c.isalnum() or c in (' ', '-', '_')).rstrip()
            output_filename = f"selected_candidates_{safe_job_title}_{timestamp}.docx"
        
        # Create Word document
        doc = Document()
        
        # Add title
        title = doc.add_heading(f'Selected Candidates - {job_title}', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add generation info
        info_paragraph = doc.add_paragraph()
        info_paragraph.add_run(f"Generated on: {datetime.now().strftime('%B %d, %Y at %I:%M %p')}\n")
        info_paragraph.add_run(f"Total Candidates: {len(selected_candidates)}")
        doc.add_paragraph()
        
        # Add candidates
        for i, candidate in enumerate(selected_candidates, 1):
            self._add_candidate_details(doc, candidate, i)
        
        # Save document
        doc.save(output_filename)
        logger.info(f"Selected candidates document saved as: {output_filename}")
        
        return output_filename
    
    def _add_document_header(self, doc):
        """Add header section to the document"""
        # Main title
        title = doc.add_heading('HR Automation - Candidate Shortlists', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Subtitle
        subtitle = doc.add_paragraph()
        subtitle_run = subtitle.add_run('Automated Candidate Matching Results')
        subtitle_run.font.size = Pt(14)
        subtitle_run.bold = True
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Generation info
        info_paragraph = doc.add_paragraph()
        info_paragraph.add_run(f"Generated on: {datetime.now().strftime('%B %d, %Y at %I:%M %p')}")
        info_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add separator
        doc.add_paragraph("=" * 80)
    
    def _add_job_section(self, doc, job_title: str, candidates: List[Dict[str, Any]]):
        """Add a job section with its candidates"""
        # Job title
        job_heading = doc.add_heading(f'{job_title}', 1)
        job_heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
        
        # Job summary
        summary_para = doc.add_paragraph()
        summary_para.add_run(f"Candidates Found: {len(candidates)}\n")
        if candidates:
            avg_score = sum(c.get('score', 0) for c in candidates) / len(candidates)
            summary_para.add_run(f"Average Match Score: {avg_score:.2f}")
        
        doc.add_paragraph()
        
        # Add each candidate
        if not candidates:
            no_candidates_para = doc.add_paragraph("No matching candidates found for this position.")
            no_candidates_para.italic = True
        else:
            for i, candidate_match in enumerate(candidates, 1):
                candidate = candidate_match.get('candidate', {})
                score = candidate_match.get('score', 0)
                matched_skills = candidate_match.get('matched_skills', [])
                
                # Convert candidate object to dict if needed
                if hasattr(candidate, 'to_dict'):
                    candidate_dict = candidate.to_dict()
                else:
                    candidate_dict = candidate
                
                self._add_candidate_details(doc, candidate_dict, i, score, matched_skills)
        
        # Add separator
        doc.add_page_break()
    
    def _add_candidate_details(self, doc, candidate: Dict[str, Any], 
                             index: int, score: float = None, matched_skills: List[str] = None):
        """Add detailed candidate information"""
        # Candidate header
        candidate_heading = doc.add_heading(f'{index}. {candidate.get("full_name", "Unknown Name")}', 2)
        
        # Create a table for candidate details
        table = doc.add_table(rows=0, cols=2)
        table.style = 'Table Grid'
        
        # Add candidate information
        details = [
            ("Full Name", candidate.get("full_name", "N/A")),
            ("LinkedIn Profile", candidate.get("linkedin_url", "N/A")),
            ("Email Address", candidate.get("email", "Not Available")),
            ("Current Position", candidate.get("position", "N/A")),
            ("Current Company", candidate.get("company", "N/A")),
            ("Connected On", candidate.get("connected_on", "N/A"))
        ]
        
        if score is not None:
            details.append(("Match Score", f"{score:.2f}"))
        
        if matched_skills:
            details.append(("Matched Skills", ", ".join(matched_skills)))
        
        if candidate.get("location"):
            details.append(("Location", candidate["location"]))
        
        if candidate.get("extracted_skills"):
            details.append(("Extracted Skills", candidate["extracted_skills"]))
        
        if candidate.get("experience_summary"):
            details.append(("Experience Summary", candidate["experience_summary"]))
        
        # Add rows to table
        for label, value in details:
            row_cells = table.add_row().cells
            row_cells[0].text = label
            row_cells[0].paragraphs[0].runs[0].bold = True
            row_cells[1].text = str(value)
        
        # Add spacing
        doc.add_paragraph()
    
    def _add_document_footer(self, doc):
        """Add footer to the document"""
        doc.add_paragraph("=" * 80)
        
        footer = doc.add_paragraph()
        footer.add_run("Document generated by HR Automation System\n")
        footer.add_run("For internal use only - Confidential candidate information")
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Make footer italic and smaller
        for run in footer.runs:
            run.italic = True
            run.font.size = Pt(10)
    
    def _create_text_document(self, shortlists: Dict[str, List[Dict[str, Any]]], 
                            output_filename: str = None) -> str:
        """Fallback method to create text document when python-docx is not available"""
        if not output_filename:
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            output_filename = f"candidate_shortlists_{timestamp}.txt"
        
        with open(output_filename, 'w', encoding='utf-8') as f:
            f.write("HR AUTOMATION - CANDIDATE SHORTLISTS\n")
            f.write("=" * 50 + "\n")
            f.write(f"Generated on: {datetime.now().strftime('%B %d, %Y at %I:%M %p')}\n\n")
            
            for job_title, candidates in shortlists.items():
                f.write(f"\nJOB: {job_title}\n")
                f.write("-" * 40 + "\n")
                f.write(f"Candidates Found: {len(candidates)}\n\n")
                
                if not candidates:
                    f.write("No matching candidates found for this position.\n\n")
                    continue
                
                for i, candidate_match in enumerate(candidates, 1):
                    candidate = candidate_match.get('candidate', {})
                    score = candidate_match.get('score', 0)
                    matched_skills = candidate_match.get('matched_skills', [])
                    
                    # Convert candidate object to dict if needed
                    if hasattr(candidate, 'to_dict'):
                        candidate_dict = candidate.to_dict()
                    else:
                        candidate_dict = candidate
                    
                    f.write(f"{i}. {candidate_dict.get('full_name', 'Unknown Name')}\n")
                    f.write(f"   LinkedIn: {candidate_dict.get('linkedin_url', 'N/A')}\n")
                    f.write(f"   Email: {candidate_dict.get('email', 'Not Available')}\n")
                    f.write(f"   Position: {candidate_dict.get('position', 'N/A')}\n")
                    f.write(f"   Company: {candidate_dict.get('company', 'N/A')}\n")
                    f.write(f"   Match Score: {score:.2f}\n")
                    f.write(f"   Matched Skills: {', '.join(matched_skills)}\n")
                    f.write("\n")
                
                f.write("\n" + "=" * 50 + "\n")
        
        logger.info(f"Text document saved as: {output_filename}")
        return output_filename
    
    def _create_text_document_selected(self, selected_candidates: List[Dict[str, Any]], 
                                     job_title: str, output_filename: str = None) -> str:
        """Create text document for selected candidates"""
        if not output_filename:
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            safe_job_title = "".join(c for c in job_title if c.isalnum() or c in (' ', '-', '_')).rstrip()
            output_filename = f"selected_candidates_{safe_job_title}_{timestamp}.txt"
        
        with open(output_filename, 'w', encoding='utf-8') as f:
            f.write(f"SELECTED CANDIDATES - {job_title}\n")
            f.write("=" * 50 + "\n")
            f.write(f"Generated on: {datetime.now().strftime('%B %d, %Y at %I:%M %p')}\n")
            f.write(f"Total Candidates: {len(selected_candidates)}\n\n")
            
            for i, candidate in enumerate(selected_candidates, 1):
                f.write(f"{i}. {candidate.get('full_name', 'Unknown Name')}\n")
                f.write(f"   LinkedIn: {candidate.get('linkedin_url', 'N/A')}\n")
                f.write(f"   Email: {candidate.get('email', 'Not Available')}\n")
                f.write(f"   Position: {candidate.get('position', 'N/A')}\n")
                f.write(f"   Company: {candidate.get('company', 'N/A')}\n")
                if candidate.get('score'):
                    f.write(f"   Match Score: {candidate['score']:.2f}\n")
                f.write("\n")
        
        logger.info(f"Selected candidates text document saved as: {output_filename}")
        return output_filename
    
    def generate_candidate_summary(self, candidate_match: Dict[str, Any], 
                                 job_title: str, candidate_name: str = None) -> str:
        """
        Generate a Word document summary for a single candidate
        
        Args:
            candidate_match: Dictionary containing candidate data and match info
            job_title: Job title the candidate is being considered for
            candidate_name: Optional candidate name for filename
            
        Returns:
            Path to the generated document
        """
        candidate = candidate_match.get('candidate', {})
        if not candidate_name:
            candidate_name = candidate.get('full_name', 'Unknown')
            candidate_name = candidate_name.replace(' ', '_')
        
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        safe_job_title = "".join(c for c in job_title if c.isalnum() or c in (' ', '-', '_')).rstrip()
        output_filename = f"{candidate_name}_{safe_job_title}_{timestamp}.docx"
        
        # Create document with single candidate
        return self.create_selected_candidates_document([candidate], job_title, output_filename)
    
    def generate_text_summary(self, candidate_match: Dict[str, Any], 
                            job_title: str, filename: str) -> str:
        """
        Generate a text file summary for a single candidate
        
        Args:
            candidate_match: Dictionary containing candidate data and match info
            job_title: Job title the candidate is being considered for
            filename: Output filename
            
        Returns:
            Path to the generated file
        """
        candidate = candidate_match.get('candidate', {})
        
        with open(filename, 'w', encoding='utf-8') as f:
            f.write(f"CANDIDATE SUMMARY - {job_title}\n")
            f.write("=" * 50 + "\n")
            f.write(f"Generated on: {datetime.now().strftime('%B %d, %Y at %I:%M %p')}\n\n")
            
            f.write(f"Name: {candidate.get('full_name', 'Unknown Name')}\n")
            f.write(f"LinkedIn: {candidate.get('linkedin_url', 'N/A')}\n")
            f.write(f"Email: {candidate.get('email', 'Not Available')}\n")
            f.write(f"Position: {candidate.get('position', 'N/A')}\n")
            f.write(f"Company: {candidate.get('company', 'N/A')}\n")
            f.write(f"Connected On: {candidate.get('connected_on', 'N/A')}\n")
            
            if candidate_match.get('score'):
                f.write(f"Match Score: {candidate_match['score']:.2f}\n")
            
            if candidate_match.get('matched_skills'):
                f.write(f"Matched Skills: {', '.join(candidate_match['matched_skills'])}\n")
            
            if candidate.get('location'):
                f.write(f"Location: {candidate['location']}\n")
            
            if candidate.get('extracted_skills'):
                f.write(f"Extracted Skills: {candidate['extracted_skills']}\n")
            
            if candidate.get('experience_summary'):
                f.write(f"Experience Summary: {candidate['experience_summary']}\n")
            
            f.write("\n" + "=" * 50 + "\n")
        
        logger.info(f"Text summary saved as: {filename}")
        return filename

class CandidateSelector:
    """Handles candidate selection logic"""
    
    def __init__(self):
        self.selected_candidates = []
    
    def load_shortlists(self, shortlists_file: str = "shortlists.json") -> Dict[str, List[Dict[str, Any]]]:
        """Load shortlists from JSON file"""
        try:
            with open(shortlists_file, 'r', encoding='utf-8') as f:
                return json.load(f)
        except FileNotFoundError:
            logger.error(f"Shortlists file not found: {shortlists_file}")
            return {}
        except Exception as e:
            logger.error(f"Error loading shortlists: {e}")
            return {}
    
    def display_candidates_for_selection(self, shortlists: Dict[str, List[Dict[str, Any]]]) -> Dict[str, List[Dict[str, Any]]]:
        """
        Display candidates and allow selection
        
        Args:
            shortlists: Dictionary of job titles and candidates
            
        Returns:
            Dictionary of selected candidates per job
        """
        selected_by_job = {}
        
        print("\n" + "="*60)
        print("CANDIDATE SELECTION INTERFACE")
        print("="*60)
        
        for job_title, candidates in shortlists.items():
            print(f"\nJob: {job_title}")
            print(f"Available candidates: {len(candidates)}")
            print("-" * 40)
            
            if not candidates:
                print("No candidates available for this job.")
                continue
            
            # Display candidates
            for i, candidate_match in enumerate(candidates, 1):
                candidate = candidate_match['candidate']
                score = candidate_match['score']
                print(f"{i}. {candidate['full_name']} - Score: {score:.2f}")
                print(f"   Position: {candidate['position']} at {candidate['company']}")
                print(f"   Email: {candidate['email']}")
                print(f"   LinkedIn: {candidate['linkedin_url']}")
                print()
            
            # Get selection
            selected_indices = self._get_candidate_selection(len(candidates))
            
            # Store selected candidates
            selected_candidates = []
            for idx in selected_indices:
                candidate_data = candidates[idx - 1]['candidate'].copy()
                candidate_data['score'] = candidates[idx - 1]['score']
                candidate_data['matched_skills'] = candidates[idx - 1]['matched_skills']
                selected_candidates.append(candidate_data)
            
            if selected_candidates:
                selected_by_job[job_title] = selected_candidates
                print(f"Selected {len(selected_candidates)} candidates for {job_title}")
            
            print("\n" + "="*60)
        
        return selected_by_job
    
    def _get_candidate_selection(self, total_candidates: int) -> List[int]:
        """Get user input for candidate selection"""
        while True:
            try:
                selection_input = input(f"\nEnter candidate numbers to select (1-{total_candidates}), separated by commas (or 'skip' to skip this job): ").strip()
                
                if selection_input.lower() == 'skip':
                    return []
                
                if not selection_input:
                    return []
                
                # Parse selections
                selections = []
                for num_str in selection_input.split(','):
                    num = int(num_str.strip())
                    if 1 <= num <= total_candidates:
                        selections.append(num)
                    else:
                        print(f"Invalid number: {num}. Must be between 1 and {total_candidates}")
                        raise ValueError("Invalid selection")
                
                return list(set(selections))  # Remove duplicates
                
            except ValueError:
                print("Invalid input. Please enter numbers separated by commas.")
                continue
            except KeyboardInterrupt:
                print("\nSelection cancelled.")
                return []

def demo_document_generation():
    """Demonstrate document generation with sample data"""
    # Sample data
    sample_shortlists = {
        "Senior Python Developer": [
            {
                'candidate': {
                    'full_name': 'John Doe',
                    'linkedin_url': 'https://linkedin.com/in/john-doe',
                    'email': 'john.doe@email.com',
                    'position': 'Senior Software Engineer',
                    'company': 'Tech Corp',
                    'connected_on': '2024-01-15'
                },
                'score': 0.85,
                'matched_skills': ['Python', 'Django', 'AWS'],
                'is_match': True
            }
        ],
        "Frontend React Developer": [
            {
                'candidate': {
                    'full_name': 'Jane Smith',
                    'linkedin_url': 'https://linkedin.com/in/jane-smith',
                    'email': 'jane.smith@email.com',
                    'position': 'Frontend Developer',
                    'company': 'StartupXYZ',
                    'connected_on': '2024-02-01'
                },
                'score': 0.78,
                'matched_skills': ['React', 'JavaScript', 'TypeScript'],
                'is_match': True
            }
        ]
    }
    
    # Generate document
    generator = CandidateDocumentGenerator()
    
    print("Generating sample shortlist document...")
    doc_path = generator.create_shortlist_document(sample_shortlists)
    print(f"Document generated: {doc_path}")
    
    # Generate selected candidates document
    selected_candidates = [sample_shortlists["Senior Python Developer"][0]['candidate']]
    selected_doc_path = generator.create_selected_candidates_document(
        selected_candidates, 
        "Senior Python Developer"
    )
    print(f"Selected candidates document generated: {selected_doc_path}")

def main():
    """Main function for testing document generation"""
    try:
        # Check if shortlists.json exists
        if Path("shortlists.json").exists():
            # Load and process real data
            selector = CandidateSelector()
            shortlists = selector.load_shortlists()
            
            if shortlists:
                print("Shortlists loaded successfully!")
                
                # Allow user to select candidates
                selected_by_job = selector.display_candidates_for_selection(shortlists)
                
                # Generate documents
                generator = CandidateDocumentGenerator()
                
                # Generate full shortlist document
                full_doc = generator.create_shortlist_document(shortlists)
                print(f"\nFull shortlist document created: {full_doc}")
                
                # Generate documents for selected candidates
                for job_title, selected_candidates in selected_by_job.items():
                    if selected_candidates:
                        selected_doc = generator.create_selected_candidates_document(
                            selected_candidates, job_title
                        )
                        print(f"Selected candidates document for '{job_title}': {selected_doc}")
            else:
                print("No shortlists found. Running demo instead...")
                demo_document_generation()
        else:
            print("No shortlists.json found. Running demo...")
            demo_document_generation()
            
    except Exception as e:
        logger.error(f"Error in main execution: {e}")
        print("Running demo due to error...")
        demo_document_generation()

if __name__ == "__main__":
    main()
